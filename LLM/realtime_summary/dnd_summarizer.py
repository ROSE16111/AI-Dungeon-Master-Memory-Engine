"""
D&D Real-time Voice Recording and Summarization System
Complete standalone implementation with advanced VAD
"""

import os
# Fix OpenMP library conflict
os.environ['KMP_DUPLICATE_LIB_OK'] = 'TRUE'

import sys
import time
import queue
import re
import threading
from typing import List, Dict
import numpy as np
import sounddevice as sd
import webrtcvad
from faster_whisper import WhisperModel
from datetime import datetime


class DnDVoiceSummarizer:
    """Complete D&D voice recording and summarization system with VAD"""

    def __init__(self, llm_model_path: str, words_per_chunk=150,
                 whisper_model_size="small", whisper_device="cpu",
                 llm_gpu_layers=0, llm_threads=8):
        """
        Initialize the complete system

        Args:
            llm_model_path: Path to GGUF model file for summarization
            words_per_chunk: Word threshold for generating summaries
            whisper_model_size: Whisper model (tiny/base/small/medium/large)
            whisper_device: Device for Whisper (cpu/cuda)
            llm_gpu_layers: GPU layers for LLM (0 for CPU, -1 for all GPU)
            llm_threads: CPU threads for LLM
        """
        print("="*70)
        print("D&D REAL-TIME VOICE RECORDING & SUMMARIZATION SYSTEM")
        print("="*70)

        # Configuration
        self.llm_model_path = llm_model_path
        self.words_per_chunk = words_per_chunk
        self.whisper_model_size = whisper_model_size
        self.whisper_device = whisper_device
        self.llm_gpu_layers = llm_gpu_layers
        self.llm_threads = llm_threads

        # Audio configuration
        self.SAMPLE_RATE = 16000
        self.CHANNELS = 1
        self.FRAME_MS = 20
        self.FRAME_SIZE = self.SAMPLE_RATE * self.FRAME_MS // 1000
        self.SILENCE_END_MS = 600
        self.PARTIAL_INTERVAL = 0.9
        self.OVERLAP_SEC = 0.2

        # Models
        self.llm = None
        self.whisper_model = None
        self.vad = None

        # Recording state
        self.is_recording = False
        self.audio_queue = queue.Queue(maxsize=200)
        self.session_start = None

        # Transcription state
        self.all_segments = []
        self.full_transcript = []
        self.transcription_buffer = ""

        # Summary state
        self.summaries = []
        self.context_buffer = []

        # Load models
        self._load_models()

    def _load_models(self):
        """Load all required models"""
        print("\n📦 Loading models...\n")

        # Load LLM for summarization
        try:
            from llama_cpp import Llama
            print(f"🔄 Loading LLM model: {self.llm_model_path}")
            self.llm = Llama(
                model_path=self.llm_model_path,
                n_ctx=2048,
                n_threads=self.llm_threads,
                n_gpu_layers=self.llm_gpu_layers,
                verbose=False
            )
            print("✅ LLM model loaded successfully!")
        except ImportError:
            print("❌ Error: llama-cpp-python not installed")
            print("   Run: pip install llama-cpp-python")
            raise
        except Exception as e:
            print(f"❌ Failed to load LLM: {str(e)}")
            raise

        # Load Whisper for transcription
        try:
            print(f"\n🔄 Loading Whisper model ({self.whisper_model_size})...")
            self.whisper_model = WhisperModel(
                self.whisper_model_size,
                device=self.whisper_device,
                compute_type="int8"
            )
            print("✅ Whisper model loaded successfully!")
        except ImportError:
            print("❌ Error: faster-whisper not installed")
            print("   Run: pip install faster-whisper")
            raise
        except Exception as e:
            print(f"❌ Failed to load Whisper: {str(e)}")
            raise

        # Initialize VAD
        try:
            self.vad = webrtcvad.Vad(2)  # Aggressiveness: 0-3
            print("✅ VAD initialized successfully!")
        except ImportError:
            print("❌ Error: webrtcvad not installed")
            print("   Run: pip install webrtcvad")
            raise

        print("\n" + "="*70)
        print("✨ All models loaded! Ready to start recording.")
        print("="*70 + "\n")

    def _audio_callback(self, indata, frames, time_info, status):
        """Audio input callback"""
        if status:
            pass  # Could log status for debugging
        self.audio_queue.put(indata.copy())

    def _is_speech_int16(self, frames_int16: np.ndarray) -> bool:
        """Check if audio frame contains speech using VAD"""
        try:
            return self.vad.is_speech(frames_int16.tobytes(), self.SAMPLE_RATE)
        except Exception:
            return False

    def _transcribe_float32(self, wave_f32: np.ndarray) -> str:
        """Transcribe audio using Whisper"""
        segments, _ = self.whisper_model.transcribe(
            wave_f32,
            language="en",
            beam_size=1,
            temperature=0.0,
            vad_filter=False,
            no_speech_threshold=0.4,
            compression_ratio_threshold=2.4,
        )
        return "".join(seg.text for seg in segments).strip()

    def _count_words(self, text: str) -> int:
        """Count words in text"""
        return len(re.findall(r'\b[a-zA-Z]+\b', text))

    def _generate_summary(self, text: str, index: int) -> str:
        """Generate summary for text"""
        context = ""
        if self.context_buffer:
            context = "Previous events:\n" + "\n".join(self.context_buffer[-3:]) + "\n\n"

        prompt = f"""<|begin_of_text|><|start_header_id|>system<|end_header_id|>

You are a D&D game session recorder. Your task is to summarize game segments into concise records.<|eot_id|><|start_header_id|>user<|end_header_id|>

{context}Current game segment:
{text}

Summarize this segment in ONE sentence following this format:
[Character name] did [specific action] and gained/discovered/encountered [result/outcome]

Requirements:
1. ONE sentence only, maximum 50 words
2. Must include: subject (who), action (what they did), outcome (result/consequence)
3. If multiple characters, focus on the main character
4. Focus on key actions, ignore dialogue details
5. Use past tense
6. Do NOT add prefixes like "Summary:" - output the summary directly

Examples:
- Elara investigated the abandoned temple's basement and discovered an enchanted elven short sword
- The dwarf warrior Grim dueled the orc chieftarch and successfully defeated him, earning the tribe's respect
- The rogue Roger failed to disarm the trap and triggered a poison needle mechanism, taking 6 points of damage<|eot_id|><|start_header_id|>assistant<|end_header_id|>

"""

        try:
            output = self.llm(
                prompt,
                max_tokens=100,
                temperature=0.2,
                top_p=0.9,
                stop=["<|eot_id|>", "\n\n", "Examples:", "Requirements:"],
                echo=False
            )

            summary = output['choices'][0]['text'].strip()
            summary = re.sub(r'^(Summary:|Note:|Recap:)\s*', '', summary, flags=re.IGNORECASE)
            summary = summary.split('\n')[0]

            return summary
        except Exception as e:
            return f"[Summary generation failed: {str(e)}]"

    def _summarization_thread(self):
        """Background thread for generating summaries"""
        last_check = time.time()

        while self.is_recording:
            time.sleep(5)  # Check every 5 seconds

            # Check if we have enough words to summarize
            word_count = self._count_words(self.transcription_buffer)

            if word_count >= self.words_per_chunk:
                print(f"\n{'='*70}")
                print(f"⏳ Generating summary ({word_count} words)...")
                print(f"{'='*70}")

                summary = self._generate_summary(
                    self.transcription_buffer,
                    len(self.summaries) + 1
                )

                print(f"📝 SUMMARY #{len(self.summaries) + 1}:")
                print(f"   {summary}")
                print(f"{'='*70}\n")

                self.summaries.append(summary)
                self.context_buffer.append(f"{len(self.summaries)}. {summary}")
                if len(self.context_buffer) > 3:
                    self.context_buffer.pop(0)

                # Clear buffer after summarizing
                self.transcription_buffer = ""

    def start_recording(self):
        """Start real-time voice recording and summarization"""
        self.is_recording = True
        self.session_start = time.time()

        print("\n" + "="*70)
        print("🎤 STARTING VOICE RECORDING")
        print("="*70)
        print(f"📊 Configuration:")
        print(f"   - Sample Rate: {self.SAMPLE_RATE} Hz")
        print(f"   - Frame Size: {self.FRAME_MS}ms")
        print(f"   - Summary Threshold: {self.words_per_chunk} words")
        print(f"   - Whisper Model: {self.whisper_model_size}")
        print("\n📢 Speak now! Press Ctrl+C to stop and save.\n")
        print("="*70 + "\n")

        # Start audio stream
        try:
            stream = sd.InputStream(
                samplerate=self.SAMPLE_RATE,
                blocksize=self.FRAME_SIZE,
                channels=self.CHANNELS,
                dtype="int16",
                callback=self._audio_callback,
            )
            stream.start()
        except Exception as e:
            print(f"❌ Failed to start audio stream: {e}")
            print("Tip: If device doesn't support 16kHz, try 48kHz and resample")
            raise

        # Start summarization thread
        summary_thread = threading.Thread(target=self._summarization_thread, daemon=True)
        summary_thread.start()

        # Main transcription loop
        speaking = False
        last_voice_time = 0.0
        last_partial_time = 0.0
        last_partial_text = ""

        tail = np.zeros(int(self.OVERLAP_SEC * self.SAMPLE_RATE), dtype=np.int16)
        buf = []
        utt_start_wall = None

        try:
            while True:
                indata = self.audio_queue.get(timeout=1)

                # Convert to int16
                if indata.ndim > 1:
                    pcm16 = np.frombuffer(indata, dtype=np.int16).reshape(-1, self.CHANNELS)[:, 0]
                else:
                    pcm16 = np.frombuffer(indata, dtype=np.int16)

                # VAD check
                voiced = self._is_speech_int16(pcm16)

                if voiced:
                    if not speaking:
                        speaking = True
                        utt_start_wall = time.time()
                    last_voice_time = time.time()
                    buf.append(pcm16)

                    # Partial transcription
                    if time.time() - last_partial_time >= self.PARTIAL_INTERVAL:
                        chunk = np.concatenate([tail, *buf]) if buf else tail
                        wave = chunk.astype(np.float32) / 32768.0
                        try:
                            text = self._transcribe_float32(wave)
                            if text and text != last_partial_text:
                                print(f"\r[partial] {text}", end="", flush=True)
                                last_partial_text = text
                        except Exception as e:
                            print(f"\n[warn] partial failed: {e}")
                        last_partial_time = time.time()

                else:
                    # Check for end of utterance
                    if speaking and (time.time() - last_voice_time) * 1000 >= self.SILENCE_END_MS:
                        speaking = False
                        utt_end_wall = time.time()
                        utter = np.concatenate([tail, *buf]) if buf else tail
                        wave = utter.astype(np.float32) / 32768.0

                        try:
                            final_text = self._transcribe_float32(wave)
                            if final_text:
                                print("\r" + " " * 120, end="\r")
                                timestamp = time.strftime("%H:%M:%S")
                                print(f"🗣️  [{timestamp}] {final_text}", flush=True)

                                # Add to transcription buffer for summarization
                                self.transcription_buffer += " " + final_text

                                # Record segment
                                rel_start = max(0.0, (utt_start_wall or time.time()) - self.session_start)
                                rel_end = max(rel_start, utt_end_wall - self.session_start)
                                self.all_segments.append({
                                    "start": rel_start,
                                    "end": rel_end,
                                    "text": final_text
                                })
                                self.full_transcript.append(final_text)
                        except Exception as e:
                            print(f"\n[warn] final transcription failed: {e}")

                        # Update tail and clear buffer
                        if utter.size >= tail.size:
                            tail = utter[-tail.size:].copy()
                        else:
                            z = np.zeros(tail.size - utter.size, dtype=np.int16)
                            tail = np.concatenate([z, utter])
                        buf.clear()
                        last_partial_text = ""
                        last_partial_time = time.time()
                        utt_start_wall = None

        except KeyboardInterrupt:
            print("\n\n" + "="*70)
            print("🛑 STOPPING RECORDING...")
            print("="*70)
        except queue.Empty:
            pass
        finally:
            # Finalize last utterance if still speaking
            if speaking and buf:
                try:
                    utter = np.concatenate([tail, *buf])
                    wave = utter.astype(np.float32) / 32768.0
                    final_text = self._transcribe_float32(wave)
                    if final_text:
                        print("\r" + " " * 120, end="\r")
                        print(f"🗣️  [final] {final_text}", flush=True)
                        self.transcription_buffer += " " + final_text
                        rel_start = max(0.0, (utt_start_wall or time.time()) - self.session_start)
                        rel_end = max(rel_start, time.time() - self.session_start)
                        self.all_segments.append({
                            "start": rel_start,
                            "end": rel_end,
                            "text": final_text
                        })
                        self.full_transcript.append(final_text)
                except Exception as e:
                    print(f"[warn] finalize failed: {e}")

            # Generate final summary if buffer has content
            if self.transcription_buffer.strip() and self._count_words(self.transcription_buffer) > 20:
                print("\n⏳ Generating final summary...")
                summary = self._generate_summary(
                    self.transcription_buffer,
                    len(self.summaries) + 1
                )
                self.summaries.append(summary)
                print(f"📝 Final summary: {summary}\n")

            # Stop stream
            self.is_recording = False
            try:
                stream.stop()
                stream.close()
            except Exception:
                pass

            # Save all outputs
            if self.full_transcript:
                self._save_all_outputs()

            print("\n[Bye]")

    def _save_all_outputs(self):
        """Save transcripts and summaries to files"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        basename = f"dnd_session_{timestamp}"

        print("\n" + "="*70)
        print("💾 SAVING FILES...")
        print("="*70)

        # Save full transcript (TXT)
        txt_path = f"{basename}_transcript.txt"
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write("="*70 + "\n")
            f.write("D&D SESSION FULL TRANSCRIPT\n")
            f.write(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("="*70 + "\n\n")
            f.write("\n\n".join(self.full_transcript).strip() + "\n")
        print(f"✅ Transcript saved: {txt_path}")

        # Save summaries
        summary_path = f"{basename}_summaries.txt"
        with open(summary_path, "w", encoding="utf-8") as f:
            f.write("="*70 + "\n")
            f.write("D&D GAME SESSION SUMMARIES\n")
            f.write(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"Total Summaries: {len(self.summaries)}\n")
            f.write("="*70 + "\n\n")
            for i, summary in enumerate(self.summaries, 1):
                f.write(f"{i}. {summary}\n")
        print(f"✅ Summaries saved: {summary_path}")

        # Save SRT (subtitle format)
        srt_path = f"{basename}_transcript.srt"
        with open(srt_path, "w", encoding="utf-8") as f:
            for i, seg in enumerate(self.all_segments, 1):
                f.write(f"{i}\n")
                f.write(seg["text"].strip() + "\n\n")
        print(f"✅ SRT saved: {srt_path}")

        # Try to save DOCX
        try:
            from docx import Document
            from docx.shared import Pt
            doc = Document()
            doc.add_heading('D&D Session Transcript', level=1)

            # Add summaries section
            doc.add_heading('Session Summaries', level=2)
            for i, summary in enumerate(self.summaries, 1):
                doc.add_paragraph(f"{i}. {summary}")

            doc.add_paragraph("")
            doc.add_heading('Full Transcript', level=2)

            for seg in self.all_segments:
                p = doc.add_paragraph(seg["text"].strip())
                p.style.font.size = Pt(12)
                doc.add_paragraph("")

            docx_path = f"{basename}_full.docx"
            doc.save(docx_path)
            print(f"✅ DOCX saved: {docx_path}")
        except Exception:
            pass  # Skip if python-docx not installed

        # Print statistics
        print("\n" + "="*70)
        print("📊 SESSION STATISTICS")
        print("="*70)
        print(f"Total Summaries: {len(self.summaries)}")
        print(f"Total Words: {sum(self._count_words(t) for t in self.full_transcript)}")
        print(f"Total Segments: {len(self.all_segments)}")
        print("="*70)

        if self.summaries:
            print("\n📖 COMPLETE SESSION SUMMARY:")
            print("="*70)
            for i, summary in enumerate(self.summaries, 1):
                print(f"{i}. {summary}")
            print("="*70)


# Main execution
if __name__ == "__main__":
    # Configuration
    LLM_MODEL_PATH = r"D:\Program Files\JetBrains\PyCharm 2025.2\Summary\llama.cpp\llama-3.1-8b-instruct-q4_K_M.gguf"

    # Settings
    WORDS_PER_SUMMARY = 150        # Generate summary every N words
    WHISPER_MODEL = "small"        # tiny/base/small/medium/large
    WHISPER_DEVICE = "cpu"         # cpu or cuda
    LLM_GPU_LAYERS = 0             # 0 for CPU, -1 for all GPU
    LLM_THREADS = 8                # CPU threads

    try:
        # Initialize system
        summarizer = DnDVoiceSummarizer(
            llm_model_path=LLM_MODEL_PATH,
            words_per_chunk=WORDS_PER_SUMMARY,
            whisper_model_size=WHISPER_MODEL,
            whisper_device=WHISPER_DEVICE,
            llm_gpu_layers=LLM_GPU_LAYERS,
            llm_threads=LLM_THREADS
        )

        # Start recording
        summarizer.start_recording()

    except Exception as e:
        print(f"\n❌ Fatal error: {str(e)}")
        import traceback
        traceback.print_exc()